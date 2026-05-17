#!/usr/bin/env python3
"""
AgileDev Suite - Script de Ingesta de Proyectos
Convierte archivos existentes (audio, imagenes, documentos) en
documentos markdown estructurados para que el agente AgileDev PM
pueda utilizarlos como guia en la preparacion del proyecto.

Uso:
    python scripts/ingesta-proyecto.py

Requisitos:
    pip install -r requirements.txt
    - Tesseract OCR: https://github.com/UB-Mannheim/tesseract/wiki
    - ffmpeg: https://ffmpeg.org/ (para Whisper)
"""

import os
import sys
from pathlib import Path
from datetime import datetime
import shutil

BASE_DIR = Path(__file__).resolve().parent.parent
PROJECTS_DIR = BASE_DIR / "proyectos"

SUPPORTED_FORMATS = {
    "audio": {".mp3", ".wav", ".m4a", ".ogg", ".wma", ".aac", ".flac"},
    "image": {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tiff", ".tif"},
    "document": {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml", ".log", ".rtf"},
}

FILE_TYPE_NAMES = {
    ".mp3": "Audio MP3", ".wav": "Audio WAV", ".m4a": "Audio M4A", ".ogg": "Audio OGG",
    ".png": "Imagen PNG", ".jpg": "Imagen JPG", ".jpeg": "Imagen JPEG", ".bmp": "Imagen BMP",
    ".webp": "Imagen WebP", ".pdf": "Documento PDF", ".docx": "Documento Word",
    ".txt": "Texto", ".md": "Markdown", ".csv": "CSV", ".json": "JSON",
}


def timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def date_only():
    return datetime.now().strftime("%Y-%m-%d")


def format_size(bytes_):
    for unit in ["B", "KB", "MB", "GB"]:
        if bytes_ < 1024:
            return f"{bytes_:.1f} {unit}"
        bytes_ /= 1024
    return f"{bytes_:.1f} TB"


def sanitize_name(name):
    """Remove or replace characters problematic for filenames."""
    invalid_chars = '<>:"/\\|?*'
    for c in invalid_chars:
        name = name.replace(c, "-")
    return name.strip(". ")[:200]


def _yaml_escape(value):
    """Escape value for YAML if it contains special characters."""
    if not value:
        return value
    if any(c in value for c in ':"#{}[]&*!|>%@`,'):
        return f'"{value}"'
    return value


def write_markdown(filepath, frontmatter, content):
    """Write a markdown file with YAML frontmatter."""
    filepath.parent.mkdir(parents=True, exist_ok=True)
    yaml_lines = ["---"]
    for k, v in frontmatter.items():
        if v is None:
            continue
        if isinstance(v, list):
            yaml_lines.append(f"{k}:")
            for item in v:
                yaml_lines.append(f"  - {_yaml_escape(str(item))}")
        elif isinstance(v, bool):
            yaml_lines.append(f"{k}: {'true' if v else 'false'}")
        elif isinstance(v, (int, float)):
            yaml_lines.append(f"{k}: {v}")
        else:
            yaml_lines.append(f"{k}: {_yaml_escape(str(v))}")
    yaml_lines.append("---")
    text = "\n".join(yaml_lines) + "\n\n" + content
    filepath.write_text(text, encoding="utf-8")
    return True


# ─── Dependency Check ────────────────────────────────────────────────────

def check_dependencies():
    """Check which processing dependencies are available."""
    deps = {}

    try:
        import whisper
        deps["whisper_local"] = True
    except ImportError:
        deps["whisper_local"] = False

    try:
        import openai
        deps["openai_api"] = True
    except ImportError:
        deps["openai_api"] = False

    try:
        from PIL import Image
        deps["pillow"] = True
    except ImportError:
        deps["pillow"] = False

    try:
        import pytesseract
        try:
            pytesseract.get_tesseract_version()
            deps["pytesseract"] = True
        except Exception:
            deps["pytesseract"] = "no_binary"
    except ImportError:
        deps["pytesseract"] = False

    try:
        import fitz
        deps["pymupdf"] = True
    except ImportError:
        deps["pymupdf"] = False

    try:
        import docx
        deps["python-docx"] = True
    except ImportError:
        deps["python-docx"] = False

    return deps


def print_dep_status(deps):
    """Display dependency status."""
    dep_names = [
        ("whisper_local", "Whisper local"),
        ("openai_api", "OpenAI API (transcripcion)"),
        ("pillow", "Pillow (imagenes)"),
        ("pytesseract", "Tesseract OCR"),
        ("pymupdf", "PyMuPDF (PDF)"),
        ("python-docx", "python-docx (Word)"),
    ]
    for key, name in dep_names:
        val = deps.get(key)
        if val is True:
            print(f"  [OK]   {name}")
        elif val == "no_binary":
            print(f"  [--]   {name} (instalado pero Tesseract OCR no encontrado en el sistema)")
        else:
            print(f"  [--]   {name}")


# ─── Project Management ─────────────────────────────────────────────────

def list_projects():
    """Return sorted list of existing project names."""
    if not PROJECTS_DIR.exists():
        return []
    return sorted([
        d.name for d in PROJECTS_DIR.iterdir()
        if d.is_dir() and not d.name.startswith(".")
    ])


def create_project(name):
    """Create a new project with the 3 input folders."""
    name = sanitize_name(name)
    if not name:
        print("  Error: nombre de proyecto invalido.")
        return None
    project_dir = PROJECTS_DIR / name
    try:
        (project_dir / "audio").mkdir(parents=True, exist_ok=True)
        (project_dir / "imagenes").mkdir(parents=True, exist_ok=True)
        (project_dir / "documentos").mkdir(parents=True, exist_ok=True)
        return project_dir
    except Exception as e:
        print(f"  Error creando proyecto: {e}")
        return None


def create_project_interactive():
    """Interactive new project creation."""
    while True:
        name = input("  Nombre del proyecto (sin espacios): ").strip()
        if not name:
            print("  El nombre no puede estar vacio.")
            continue
        name = sanitize_name(name)
        project_dir = PROJECTS_DIR / name
        if project_dir.exists():
            print(f"  El proyecto '{name}' ya existe.")
            overwrite = input("  Agregar archivos al proyecto existente? (S/n): ").strip().lower()
            if overwrite not in ("", "s", "si", "y", "yes"):
                continue
            return name
        result = create_project(name)
        if result:
            print(f"\n  Proyecto '{name}' creado en: {result}")
            print(f"  Coloca tus archivos en las carpetas:")
            print(f"    {result / 'audio'}")
            print(f"    {result / 'imagenes'}")
            print(f"    {result / 'documentos'}")
            return name
        return None


def select_project_interactive(projects):
    """Show interactive project selection menu."""
    print("\n  Proyectos disponibles:")
    for i, name in enumerate(projects, 1):
        project_dir = PROJECTS_DIR / name
        audio_count = len(list((project_dir / "audio").glob("*"))) if (project_dir / "audio").exists() else 0
        img_count = len(list((project_dir / "imagenes").glob("*"))) if (project_dir / "imagenes").exists() else 0
        doc_count = len(list((project_dir / "documentos").glob("*"))) if (project_dir / "documentos").exists() else 0
        print(f"  {i}. {name}  (audio: {audio_count}, imagenes: {img_count}, docs: {doc_count})")

    print(f"  {len(projects) + 1}. Crear nuevo proyecto")
    print("  0. Salir")

    while True:
        try:
            choice = input("\n  Opcion: ").strip()
            if choice == "0":
                return None, False
            choice = int(choice)
            if 1 <= choice <= len(projects):
                return projects[choice - 1], False
            elif choice == len(projects) + 1:
                name = create_project_interactive()
                if name:
                    return name, True
                return None, False
            else:
                print("  Opcion invalida.")
        except ValueError:
            print("  Ingresa un numero valido.")


# ─── File Processing ────────────────────────────────────────────────────

def process_audio_file(file_path, output_dir, deps):
    """Transcribe audio using local Whisper with OpenAI API fallback."""
    rel_path = file_path.relative_to(PROJECTS_DIR)
    stat = file_path.stat()
    duration_str = "desconocida"

    frontmatter = {
        "fuente": str(rel_path),
        "tipo": "transcripcion",
        "formato": FILE_TYPE_NAMES.get(file_path.suffix.lower(), "Audio"),
        "tamano": format_size(stat.st_size),
        "fecha_procesamiento": date_only(),
    }

    lines = []
    lines.append(f"# Transcripcion: {file_path.stem}")
    lines.append("")
    lines.append("## Metadatos")
    lines.append("")
    lines.append(f"- **Fuente**: {rel_path}")
    lines.append(f"- **Formato**: {FILE_TYPE_NAMES.get(file_path.suffix.lower(), 'Audio')}")
    lines.append(f"- **Tamano**: {format_size(stat.st_size)}")

    text_content = None

    # Try local Whisper
    if deps.get("whisper_local"):
        try:
            import whisper
            print(f"    Cargando modelo Whisper (base)...")
            model = whisper.load_model("base")
            print(f"    Transcribiendo audio...")
            result = model.transcribe(str(file_path), language="es")

            if "duration" in result and result["duration"]:
                m, s = divmod(int(result["duration"]), 60)
                h, m = divmod(m, 60)
                duration_str = f"{h}h {m}m {s}s" if h else f"{m}m {s}s"
                frontmatter["duracion"] = duration_str
                lines.append(f"- **Duracion**: {duration_str}")

            lines.append(f"- **Procesado**: {timestamp()}")
            lines.append("")
            lines.append("## Transcripcion")
            lines.append("")

            segments = result.get("segments", [])
            if segments:
                for seg in segments:
                    start = seg["start"]
                    h, m, s = int(start // 3600), int((start % 3600) // 60), int(start % 60)
                    text = seg["text"].strip()
                    lines.append(f"[{h:02d}:{m:02d}:{s:02d}] {text}")
            else:
                text_content = result.get("text", "").strip()
                if text_content:
                    lines.append(text_content)
                else:
                    lines.append("*No se detecto contenido de audio.*")

            return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))

        except Exception as e:
            print(f"    Error con Whisper local: {e}")
            if not deps.get("openai_api"):
                print(f"    No hay metodo alternativo disponible.")
                return False

    # Fallback to OpenAI API
    if deps.get("openai_api"):
        try:
            from openai import OpenAI
            client = OpenAI()
            print(f"    Enviando a OpenAI Whisper API...")
            with open(file_path, "rb") as f:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f,
                    response_format="verbose_json",
                    language="es",
                )

            if hasattr(transcript, "duration") and transcript.duration:
                m, s = divmod(int(transcript.duration), 60)
                h, m = divmod(m, 60)
                duration_str = f"{h}h {m}m {s}s" if h else f"{m}m {s}s"
                frontmatter["duracion"] = duration_str
                lines.append(f"- **Duracion**: {duration_str}")

            lines.append(f"- **Procesado**: {timestamp()}")
            lines.append("")
            lines.append("## Transcripcion")
            lines.append("")

            text = getattr(transcript, "text", str(transcript))
            if text and text.strip():
                lines.append(text.strip())
            else:
                lines.append("*No se detecto contenido de audio.*")

            return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))

        except Exception as e:
            print(f"    Error con OpenAI API: {e}")
            return False

    print(f"    No hay modulo de transcripcion disponible.")
    print(f"    Instala openai-whisper (local) u openai (API).")
    return False


def process_image_file(file_path, output_dir, deps):
    """Extract text from image via OCR and capture metadata."""
    rel_path = file_path.relative_to(PROJECTS_DIR)
    stat = file_path.stat()

    frontmatter = {
        "fuente": str(rel_path),
        "tipo": "imagen",
        "formato": FILE_TYPE_NAMES.get(file_path.suffix.lower(), "Imagen"),
        "tamano": format_size(stat.st_size),
        "fecha_procesamiento": date_only(),
    }

    lines = []
    lines.append(f"# Imagen: {file_path.stem}")
    lines.append("")
    lines.append("## Metadatos")
    lines.append("")
    lines.append(f"- **Fuente**: {rel_path}")
    lines.append(f"- **Formato**: {FILE_TYPE_NAMES.get(file_path.suffix.lower(), 'Imagen')}")
    lines.append(f"- **Tamano**: {format_size(stat.st_size)}")

    if deps.get("pillow"):
        try:
            from PIL import Image
            img = Image.open(file_path)
            width, height = img.size
            frontmatter["dimensiones"] = f"{width}x{height}"
            lines.append(f"- **Dimensiones**: {width}x{height}")
            frontmatter["modo_color"] = img.mode
            lines.append(f"- **Modo color**: {img.mode}")
            img.close()
        except Exception as e:
            print(f"    Error leyendo metadatos de imagen: {e}")

    lines.append(f"- **Procesado**: {timestamp()}")

    if deps.get("pytesseract") and deps.get("pillow"):
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="spa")
            img.close()

            if text and text.strip():
                lines.append("")
                lines.append("## Texto extraido (OCR)")
                lines.append("")
                lines.append("```")
                lines.append(text.strip())
                lines.append("```")
            else:
                lines.append("")
                lines.append("*No se detecto texto en la imagen.*")
        except Exception as e:
            print(f"    Error en OCR: {e}")
            lines.append("")
            lines.append(f"*Error al procesar OCR: {e}*")
    else:
        missing = []
        if not deps.get("pillow"):
            missing.append("Pillow")
        if not deps.get("pytesseract"):
            missing.append("pytesseract + Tesseract OCR")
        lines.append("")
        lines.append(f"*OCR no disponible. Instala: {', '.join(missing)}*")

    return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))


def process_pdf_file(file_path, output_dir, deps):
    """Extract text from PDF."""
    rel_path = file_path.relative_to(PROJECTS_DIR)
    stat = file_path.stat()

    frontmatter = {
        "fuente": str(rel_path),
        "tipo": "pdf",
        "tamano": format_size(stat.st_size),
        "fecha_procesamiento": date_only(),
    }

    lines = []
    lines.append(f"# Documento: {file_path.stem}")
    lines.append("")
    lines.append("## Metadatos")
    lines.append("")
    lines.append(f"- **Fuente**: {rel_path}")
    lines.append(f"- **Tamano**: {format_size(stat.st_size)}")

    if deps.get("pymupdf"):
        try:
            import fitz
            doc = fitz.open(file_path)
            page_count = len(doc)
            frontmatter["paginas"] = page_count
            lines.append(f"- **Paginas**: {page_count}")

            metadata = doc.metadata
            if metadata:
                for yaml_key, pdf_key in [("autor", "author"), ("titulo", "title")]:
                    val = metadata.get(pdf_key, "").strip()
                    if val:
                        frontmatter[yaml_key] = val
                        lines.append(f"- **{yaml_key.capitalize()}**: {val}")

            lines.append(f"- **Procesado**: {timestamp()}")
            lines.append("")
            lines.append("## Contenido extraido")
            lines.append("")

            all_text = []
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                text = page.get_text().strip()
                if text:
                    all_text.append(f"### Pagina {page_num + 1}")
                    all_text.append("")
                    all_text.append(text)
                    all_text.append("")

            doc.close()

            if all_text:
                lines.extend(all_text)
            else:
                lines.append("*No se pudo extraer texto (PDF escaneado o sin contenido textual).*")
                lines.append("*Considera usar OCR externo para este archivo.*")

            return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))

        except Exception as e:
            print(f"    Error procesando PDF: {e}")
            return False

    print(f"    PyMuPDF no disponible. Instala: pip install PyMuPDF")
    lines.append(f"**Error**: No se pudo procesar el PDF (PyMuPDF no instalado).")
    lines.append("")
    lines.append("---")
    lines.append(f"### Metadatos del archivo")
    lines.append(f"- Ruta: {rel_path}")
    lines.append(f"- Tamano: {format_size(stat.st_size)}")
    return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))


def process_docx_file(file_path, output_dir, deps):
    """Extract text and tables from Word document."""
    rel_path = file_path.relative_to(PROJECTS_DIR)
    stat = file_path.stat()

    frontmatter = {
        "fuente": str(rel_path),
        "tipo": "word",
        "tamano": format_size(stat.st_size),
        "fecha_procesamiento": date_only(),
    }

    lines = []
    lines.append(f"# Documento: {file_path.stem}")
    lines.append("")
    lines.append("## Metadatos")
    lines.append("")
    lines.append(f"- **Fuente**: {rel_path}")
    lines.append(f"- **Tamano**: {format_size(stat.st_size)}")

    if deps.get("python-docx"):
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            frontmatter["parrafos"] = len(paragraphs)
            lines.append(f"- **Parrafos**: {len(paragraphs)}")
            lines.append(f"- **Procesado**: {timestamp()}")
            lines.append("")
            lines.append("## Contenido extraido")
            lines.append("")

            if paragraphs:
                lines.append("\n\n".join(paragraphs))
            else:
                lines.append("*El documento no contiene texto en parrafos.*")

            tables = doc.tables
            if tables:
                lines.append("")
                lines.append(f"## Tablas ({len(tables)})")
                lines.append("")
                for i, table in enumerate(tables):
                    lines.append(f"### Tabla {i + 1}")
                    lines.append("")
                    for row in table.rows:
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        lines.append("| " + " | ".join(cells) + " |")
                    lines.append("")

            return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))

        except Exception as e:
            print(f"    Error procesando Word: {e}")
            return False

    print(f"    python-docx no disponible. Instala: pip install python-docx")
    lines.append(f"**Error**: No se pudo procesar el documento (python-docx no instalado).")
    return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))


def process_text_file(file_path, output_dir, deps=None):
    """Read and copy plain text files."""
    rel_path = file_path.relative_to(PROJECTS_DIR)
    stat = file_path.stat()

    frontmatter = {
        "fuente": str(rel_path),
        "tipo": "texto",
        "formato": FILE_TYPE_NAMES.get(file_path.suffix.lower(), "Texto"),
        "tamano": format_size(stat.st_size),
        "fecha_procesamiento": date_only(),
    }

    lines = []
    lines.append(f"# Documento: {file_path.stem}")
    lines.append("")
    lines.append("## Metadatos")
    lines.append("")
    lines.append(f"- **Fuente**: {rel_path}")
    lines.append(f"- **Formato**: {FILE_TYPE_NAMES.get(file_path.suffix.lower(), 'Texto')}")
    lines.append(f"- **Tamano**: {format_size(stat.st_size)}")
    lines.append(f"- **Procesado**: {timestamp()}")
    lines.append("")
    lines.append("## Contenido")
    lines.append("")

    try:
        content = file_path.read_text(encoding="utf-8", errors="replace")
        content = content.strip()
        if file_path.suffix.lower() in (".md",):
            lines.append(content)
        else:
            lines.append("```" + file_path.suffix[1:] if file_path.suffix else "")
            lines.append(content)
            lines.append("```")
    except Exception as e:
        print(f"    Error leyendo archivo: {e}")
        lines.append(f"*Error al leer archivo: {e}*")

    return write_markdown(output_dir / f"{sanitize_name(file_path.stem)}.md", frontmatter, "\n".join(lines))


# ─── Index Generation ───────────────────────────────────────────────────

def generate_index(project_name, output_dir, results):
    """Generate the main index file linking all processed materials."""
    frontmatter = {
        "titulo": f"Insumos del proyecto: {project_name}",
        "tipo": "indice",
        "fecha_generacion": date_only(),
        "total_procesados": len(results["processed"]),
        "total_errores": len(results["errors"]),
        "total_omitidos": len(results["skipped"]),
    }

    lines = []
    lines.append(f"# Indice de Insumos - {project_name}")
    lines.append("")
    lines.append(f"Generado el: {timestamp()}")
    lines.append("")
    lines.append("## Resumen")
    lines.append("")
    lines.append(f"- Archivos procesados: {len(results['processed'])}")
    lines.append(f"- Errores: {len(results['errors'])}")
    lines.append(f"- Omitidos: {len(results['skipped'])}")
    lines.append("")

    folder_map = {
        "audio": "transcripciones",
        "imagenes": "imagenes",
        "documentos": "documentos",
    }

    sections = {
        "transcripciones": ("Transcripciones de Audio", []),
        "imagenes": ("Imagenes Procesadas", []),
        "documentos": ("Documentos Procesados", []),
    }

    for item in results["processed"]:
        p = Path(item)
        source_type = p.parts[0] if len(p.parts) > 0 else ""
        target = folder_map.get(source_type, "otros")
        md_name = sanitize_name(p.stem) + ".md"
        link = f"{target}/{md_name}"
        label = str(p)
        if target in sections:
            sections[target][1].append(f"  - [{label}]({link})")

    for section_key, (section_title, items) in sections.items():
        if items:
            lines.append(f"## {section_title}")
            lines.append("")
            lines.extend(items)
            lines.append("")

    if results["errors"]:
        lines.append("## Archivos con errores")
        lines.append("")
        for item in results["errors"]:
            lines.append(f"- {item}")
        lines.append("")

    if results["skipped"]:
        lines.append("## Archivos omitidos (formato no soportado)")
        lines.append("")
        for item in results["skipped"]:
            lines.append(f"- {item}")
        lines.append("")

    return write_markdown(output_dir / "00-indice.md", frontmatter, "\n".join(lines))


# ─── Project Processing ─────────────────────────────────────────────────

def scan_files(project_dir):
    """Scan input folders and classify files."""
    folders = {
        "audio": (project_dir / "audio", SUPPORTED_FORMATS["audio"], "audio"),
        "imagenes": (project_dir / "imagenes", SUPPORTED_FORMATS["image"], "imagen"),
        "documentos": (project_dir / "documentos", SUPPORTED_FORMATS["document"], "document"),
    }

    files = []
    for folder_name, (folder_path, valid_exts, file_type) in folders.items():
        if not folder_path.exists():
            continue
        for f in sorted(folder_path.iterdir()):
            if not f.is_file() or f.name.startswith("."):
                continue
            ext = f.suffix.lower()
            if ext not in valid_exts:
                files.append((f, file_type, "skip"))
            else:
                files.append((f, file_type, "process"))
    return files


def process_project(project_name, deps):
    """Process all files in a project's input folders."""
    project_dir = PROJECTS_DIR / project_name
    output_dir = project_dir / "docs" / "insumos"

    handlers = {
        ("audio", "process"): process_audio_file,
        ("imagen", "process"): process_image_file,
        ("document", "process"): None,
    }

    results = {
        "processed": [],
        "errors": [],
        "skipped": [],
    }

    files = scan_files(project_dir)

    if not files:
        print(f"\n  No se encontraron archivos en las carpetas de '{project_name}'.")
        print(f"  Coloca archivos en:")
        print(f"    {project_dir / 'audio'}")
        print(f"    {project_dir / 'imagenes'}")
        print(f"    {project_dir / 'documentos'}")
        return

    total_inputs = sum(1 for _, _, action in files if action == "process")
    total_skips = sum(1 for _, _, action in files if action == "skip")
    print(f"\n  Archivos encontrados: {total_inputs} para procesar", end="")
    if total_skips:
        print(f", {total_skips} omitidos (formato no soportado)", end="")
    print()

    for file_path, file_type, action in files:
        rel = file_path.relative_to(project_dir)
        indent = "    "

        if action == "skip":
            results["skipped"].append(str(rel))
            print(f"{indent}[--] {rel} (formato no soportado)")
            continue

        # Determine handler
        if file_type == "audio":
            handler = process_audio_file
        elif file_type == "imagen":
            handler = process_image_file
        elif file_type == "document":
            ext = file_path.suffix.lower()
            if ext == ".pdf":
                handler = process_pdf_file
            elif ext == ".docx":
                handler = process_docx_file
            else:
                handler = process_text_file
        else:
            results["skipped"].append(str(rel))
            print(f"{indent}[--] {rel} (tipo no reconocido)")
            continue

        # Determine output subfolder
        out_subdir = {
            "audio": output_dir / "transcripciones",
            "imagen": output_dir / "imagenes",
            "document": output_dir / "documentos",
        }.get(file_type, output_dir)

        try:
            print(f"{indent}Procesando: {rel}...", end="")
            sys.stdout.flush()
            result = handler(file_path, out_subdir, deps)
            if result:
                results["processed"].append(str(rel))
                print(f" [OK]")
            else:
                results["errors"].append(str(rel))
                print(f" [ERR]")
        except Exception as e:
            results["errors"].append(str(rel))
            print(f" [ERR] - {e}")

    # Generate index
    generate_index(project_name, output_dir, results)

    print(f"\n  === Resumen ===")
    print(f"  Procesados: {len(results['processed'])}")
    print(f"  Errores:    {len(results['errors'])}")
    print(f"  Omitidos:   {len(results['skipped'])}")
    print(f"\n  Documentos generados en: {output_dir}")
    print(f"  Indice: {output_dir / '00-indice.md'}")

    return results


# ─── Entry Point ────────────────────────────────────────────────────────

def main():
    print("============================================")
    print("  AgileDev Suite - Ingesta de Proyectos")
    print("============================================")
    print()

    # Check dependencies
    deps = check_dependencies()
    print("Dependencias detectadas:")
    print_dep_status(deps)
    print()

    missing_deps = [k for k, v in deps.items() if v is False]
    if missing_deps:
        print("  Nota: Algunas dependencias no estan instaladas.")
        print("  Instalacion completa: pip install -r requirements.txt")
        print()

    # Interactive menu
    projects = list_projects()

    if projects:
        print(f"Se encontraron {len(projects)} proyecto(s).")
        project_name, is_new = select_project_interactive(projects)
    else:
        print("No hay proyectos existentes.")
        print()
        create = input("Deseas crear un nuevo proyecto? (s/N): ").strip().lower()
        if create not in ("s", "si", "y", "yes"):
            print("Saliendo...")
            return
        project_name = create_project_interactive()

    if not project_name:
        print("Saliendo...")
        return

    # Process the project
    process_project(project_name, deps)

    print()
    print("Proceso completado. Los documentos generados estan listos")
    print("para ser utilizados por el agente AgileDev PM.")
    print("Abre opencode y usa Tab para cambiar al agente agile-dev-pm.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nInterrumpido por el usuario. Saliendo...")
        sys.exit(0)
    except Exception as e:
        print(f"\nError inesperado: {e}")
        print("Reporta el error en: https://github.com/anomalyco/opencode/issues")
        sys.exit(1)
